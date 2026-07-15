"""Sinh file mẫu cho test. Chạy một lần mỗi phiên, không commit file nhị phân."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest

FIXTURES = Path(__file__).parent / "fixtures"

# Times New Roman: font mặc định của hầu hết văn bản hành chính Việt Nam, và là
# font đã phơi ra lỗi glyph U+00AD (xem test_bullet_soft_hyphen).
FONT_REGULAR = Path(r"C:\Windows\Fonts\times.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\timesbd.ttf")


def _build_vietnamese_pdf(dest: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_font(fontname="TI", fontfile=str(FONT_REGULAR))
    page.insert_font(fontname="TB", fontfile=str(FONT_BOLD))

    state = {"y": 60.0}

    def put(text: str, size: float, font: str = "TI") -> None:
        page.insert_text((60, state["y"]), text, fontsize=size, fontname=font)
        state["y"] += size * 1.6

    put("Báo cáo Kỹ thuật Quý IV", 22, "TB")
    state["y"] += 8
    put("1. Tổng quan hệ thống", 15, "TB")
    put("Hệ thống được triển khai trên nền tảng đám mây, đảm bảo khả năng", 11)
    put("mở rộng linh hoạt và độ sẵn sàng cao cho toàn bộ nghiệp vụ.", 11)
    state["y"] += 12
    put("Đoạn thứ hai nói về hiệu năng và chi phí vận hành hàng tháng.", 11)
    state["y"] += 12
    put("2. Danh mục thiết bị", 15, "TB")
    put("- Máy chủ ứng dụng đặt tại Hà Nội", 11)
    put("- Thiết bị lưu trữ dự phòng ở Đà Nẵng", 11)
    state["y"] += 14
    put("2.1 Bảng chi phí", 12, "TB")

    rows = [
        ["Hạng mục", "Chi phí", "Ghi chú"],
        ["Máy chủ", "12 triệu", "Đã thanh toán"],
        ["Đường truyền", "5 triệu", "Còn nợ"],
    ]
    tx, ty, cw, rh = 60, state["y"], 145, 24
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            rect = fitz.Rect(tx + c * cw, ty + r * rh, tx + (c + 1) * cw, ty + (r + 1) * rh)
            page.draw_rect(rect, color=(0, 0, 0), width=0.7)
            page.insert_text(
                (rect.x0 + 5, rect.y0 + 16), val, fontsize=10,
                fontname="TB" if r == 0 else "TI",
            )
    doc.save(dest)
    doc.close()


def _rasterize(src: Path, dest: Path, dpi: int = 200) -> None:
    """Biến PDF thành ảnh thuần, mô phỏng bản scan không có lớp text."""
    doc = fitz.open(src)
    pix = doc[0].get_pixmap(dpi=dpi)
    out = fitz.open()
    page = out.new_page(width=doc[0].rect.width, height=doc[0].rect.height)
    page.insert_image(page.rect, pixmap=pix)
    out.save(dest)
    out.close()
    doc.close()


def _build_docx(dest: Path) -> None:
    from docx import Document

    d = Document()
    d.add_heading("Tài liệu Word mẫu", level=1)
    d.add_paragraph("Đoạn văn thường trong Word.")
    p = d.add_paragraph("Dòng này có ")
    p.add_run("chữ đậm").bold = True
    p.add_run(" và ")
    p.add_run("chữ nghiêng").italic = True
    p.add_run(" trộn lẫn.")
    d.add_heading("Mục con", level=2)
    d.add_paragraph("Mục bullet thứ nhất", style="List Bullet")
    d.add_paragraph("Mục bullet thứ hai", style="List Bullet")
    t = d.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    for c, v in zip(t.rows[0].cells, ["Cột A", "Cột B", "Cột C"]):
        c.text = v
    for c, v in zip(t.rows[1].cells, ["1", "2", "3"]):
        c.text = v
    d.save(dest)


def _build_xlsx(dest: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "DuLieu"
    for row in [["Tháng", "Doanh thu"], ["01", 1000], ["02", 1200]]:
        ws.append(row)
    wb.save(dest)


def _build_legacy_doc(dest: Path) -> bool:
    """Tạo file .doc nhị phân thật (Word 97-2003). Cần có Microsoft Word.

    Cách làm: dựng .docx bằng python-docx rồi nhờ Word lưu ngược thành .doc.
    KHÔNG dựng thẳng bằng Word COM — đã thử và mất sạch nội dung, chỉ còn mỗi
    bảng, vì Paragraphs.Add() rồi gán Range.Text hành xử khác hẳn mong đợi.
    """
    try:
        import pythoncom
        import win32com.client
        from docx import Document
    except ImportError:
        return False

    tmp_docx = dest.with_suffix(".tam.docx")
    d = Document()
    d.add_heading("Giải pháp và Phương pháp luận", level=1)
    d.add_paragraph("Đoạn mở đầu của tài liệu Word 97-2003.")
    d.add_heading("1. Phạm vi áp dụng", level=2)
    d.add_paragraph("Nội dung mô tả phạm vi triển khai hệ thống.")
    p = d.add_paragraph("Dòng có ")
    p.add_run("chữ đậm").bold = True
    p.add_run(" và ")
    p.add_run("chữ nghiêng").italic = True
    p.add_run(" trộn lẫn.")
    d.add_paragraph("Mục thứ nhất", style="List Bullet")
    d.add_paragraph("Mục thứ hai", style="List Bullet")
    t = d.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    for c, v in zip(t.rows[0].cells, ["Hạng mục", "Giá trị", "Ghi chú"]):
        c.text = v
    for c, v in zip(t.rows[1].cells, ["Máy chủ", "12 triệu", "Đã duyệt"]):
        c.text = v
    d.save(tmp_docx)

    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        try:
            doc = word.Documents.Open(str(tmp_docx.resolve()), AddToRecentFiles=False)
            doc.SaveAs2(str(dest.resolve()), FileFormat=0)  # 0 = .doc 97-2003
            doc.Close(SaveChanges=0)
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
    except Exception:  # noqa: BLE001
        return False
    finally:
        tmp_docx.unlink(missing_ok=True)
    return dest.exists()


@pytest.fixture(scope="session")
def fixtures() -> dict[str, Path]:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    paths = {
        "pdf": FIXTURES / "tiengviet.pdf",
        "scan": FIXTURES / "tiengviet_scan.pdf",
        "docx": FIXTURES / "sample.docx",
        "xlsx": FIXTURES / "sample.xlsx",
    }
    if not paths["pdf"].exists():
        _build_vietnamese_pdf(paths["pdf"])
    if not paths["scan"].exists():
        _rasterize(paths["pdf"], paths["scan"])
    if not paths["docx"].exists():
        _build_docx(paths["docx"])
    if not paths["xlsx"].exists():
        _build_xlsx(paths["xlsx"])

    # .doc cần Microsoft Word để tạo; máy nào không có thì test liên quan tự bỏ qua.
    doc = FIXTURES / "tailieu_cu.doc"
    if doc.exists() or _build_legacy_doc(doc):
        paths["doc"] = doc
    return paths
